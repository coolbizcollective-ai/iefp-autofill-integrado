import os, json
from pathlib import Path
from typing import Dict, Any, List
import pandas as pd
from docx import Document

# --- IA helpers ---
def ia_disponivel() -> bool:
    return bool(os.getenv("OPENAI_API_KEY"))

def gerar_texto_ia(titulo: str, instrucoes: str, contexto: Dict[str, Any]) -> str:
    if ia_disponivel():
        try:
            from openai import OpenAI
            client = OpenAI()
            prompt = (
                f"Escreve uma secção clara para '{titulo}' em PT-PT. "
                f"Instruções: {instrucoes}. Contexto: {json.dumps(contexto, ensure_ascii=False)}. "
                f"120-200 palavras, objetivo e profissional."
            )
            resp = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role":"user","content":prompt}],
                temperature=0.5,
            )
            return resp.choices[0].message.content.strip()
        except Exception:
            return f"[Gerar com IA] {titulo}: {instrucoes}"
    return f"[Preencher] {titulo}: {instrucoes}"

def cortar(texto: str, limite: int) -> str:
    if not texto:
        return ""
    if len(texto) <= limite:
        return texto
    recorte = texto[:limite]
    if " " in recorte:
        recorte = recorte.rsplit(" ", 1)[0]
    return recorte + "…"

# --- FINANCEIROS (Parte 2) ---
def calcular_financeiros(
    anos: List[int],
    assum: Dict[str, float],
    vendas_df: pd.DataFrame,
    pessoal_df: pd.DataFrame,
    investimento_df: pd.DataFrame,
) -> Dict[str, pd.DataFrame]:

    years = [str(a) for a in anos]

    # Vendas: calcula Y1 e projeta crescimento para restantes anos
    linhas_v = []
    for _, v in vendas_df.fillna(0).iterrows():
        y1 = float(v.get("preco",0)) * float(v.get("qtd_mensal",0)) * float(v.get("meses_y1",12))
        row = {"designacao": v.get("designacao","—"), years[0]: round(y1,2)}
        for i in range(1, len(anos)):
            prev = row[years[i-1]]
            row[years[i]] = round(prev * (1 + float(assum.get("crescimento_receitas", 0.08))), 2)
        linhas_v.append(row)
    df_vendas = pd.DataFrame(linhas_v) if linhas_v else pd.DataFrame(columns=["designacao", *years])
    tot_vendas = {y: (df_vendas[y].sum() if y in df_vendas else 0.0) for y in years}

    # COGS e FSE em percentagem das vendas
    margem = float(assum.get("margem_bruta_target", 0.55))
    cogs = {y: round((1 - margem) * tot_vendas[y], 2) for y in years}
    df_cogs = pd.DataFrame([{"rubrica": "COGS", **cogs}])

    fse_pct = float(assum.get("fse_pct_receitas", 0.12))
    fse = {y: round(fse_pct * tot_vendas[y], 2) for y in years}
    df_fse = pd.DataFrame([{"rubrica": "FSE", **fse}])

    # Pessoal + encargos sociais e aumentos anuais
    enc_soc = float(assum.get("encargos_sociais_pct", 0.2375))
    aum = float(assum.get("aumento_salarios_pct", 0.03))
    linhas_p = []
    for _, p in pessoal_df.fillna(0).iterrows():
        base = float(p.get("venc_mensal",0)) * float(p.get("n",0)) * float(p.get("meses",12))
        y_vals = {years[0]: base * (1 + enc_soc)}
        for i in range(1, len(anos)):
            prev_base = y_vals[years[i-1]] / (1 + enc_soc)
            y_vals[years[i]] = (prev_base * (1 + aum)) * (1 + enc_soc)
        linhas_p.append({"rubrica": p.get("funcao","—"), **{k: round(v,2) for k,v in y_vals.items()}})
    df_pessoal = pd.DataFrame(linhas_p) if linhas_p else pd.DataFrame(columns=["rubrica", *years])

    # Depreciações lineares por tipo
    dep_map = {
        "equipamento": int(assum.get("dep_equipamento_anos", 5)),
        "informatica": int(assum.get("dep_informatica_anos", 3)),
        "veiculos": int(assum.get("dep_veiculos_anos", 4)),
        "intangiveis": int(assum.get("dep_intangiveis_anos", 3)),
        "outros": int(assum.get("dep_outros_anos", 4)),
    }
    dep_rows, dep_tot = [], {y: 0.0 for y in years}
    for _, it in investimento_df.fillna("").iterrows():
        tipo = str(it.get("tipo","outros")).lower()
        val = float(it.get("valor", 0.0))
        anos_dep = dep_map.get(tipo, dep_map["outros"]) or 1
        anu = val / anos_dep
        row = {"bem": f"{tipo}: {it.get('descricao','')}", **{y: round(anu,2) for y in years}}
        dep_rows.append(row)
        for y in years:
            dep_tot[y] += anu
    df_dep = pd.DataFrame(dep_rows) if dep_rows else pd.DataFrame(columns=["bem", *years])

    # Empréstimo: amortização constante
    df_fin = pd.DataFrame(columns=["ano","prestacao","capital","juros","divida_final"])
    juros_tot = {y: 0.0 for y in years}
    saldo = float(assum.get("emprestimo_montante", 0.0))
    taxa = float(assum.get("emprestimo_taxa", 0.06))
    anos_amort = int(assum.get("emprestimo_anos", 3)) or 1
    amort = saldo / anos_amort if saldo > 0 else 0.0
    for i, ano in enumerate(anos):
        j = saldo * taxa if saldo > 0 else 0.0
        c = min(amort, saldo) if saldo > 0 else 0.0
        p = j + c
        saldo = max(0.0, saldo - c)
        juros_tot[str(ano)] = j
        df_fin.loc[len(df_fin)] = {
            "ano": ano, "prestacao": round(p,2), "capital": round(c,2),
            "juros": round(j,2), "divida_final": round(saldo,2)
        }

    # Demonstração de Resultados
    df_dr = pd.DataFrame([
        {"rubrica": "Vendas/Serviços", **{y: round(tot_vendas[y],2) for y in years}},
        {"rubrica": "COGS", **{y: round(df_cogs[y].iloc[0] if not df_cogs.empty else 0.0,2) for y in years}},
        {"rubrica": "FSE", **{y: round(df_fse[y].sum() if y in df_fse else 0.0,2) for y in years}},
        {"rubrica": "Pessoal", **{y: round(df_pessoal[y].sum() if not df_pessoal.empty else 0.0,2) for y in years}},
        {"rubrica": "Depreciações", **{y: round(dep_tot[y],2) for y in years}},
        {"rubrica": "Juros", **{y: round(juros_tot[y],2) for y in years}},
    ])
    res = {"rubrica": "Resultado"}
    for y in years:
        res[y] = round(
            df_dr[y].iloc[0] - df_dr[y].iloc[1] - df_dr[y].iloc[2] - df_dr[y].iloc[3] - df_dr[y].iloc[4] - df_dr[y].iloc[5],
            2
        )
    df_dr.loc[len(df_dr)] = res

    # Balanço (muito simplificado)
    inv_total = float(investimento_df.get("valor", pd.Series(dtype=float)).sum()) if not investimento_df.empty else 0.0
    df_bal = pd.DataFrame([
        {"rubrica": "Ativo Não Corrente", years[0]: round(inv_total,2)},
        {"rubrica": "Ativo Corrente", **{y: round(0.1*tot_vendas[y],2) for y in years}},
        {"rubrica": "Capital Próprio", years[0]: round(float(assum.get("capitais_proprios_iniciais",0.0)),2)},
    ])

    return {
        "vendas": df_vendas,
        "cogs": df_cogs,
        "fse": df_fse,
        "pessoal": df_pessoal,
        "depreciacoes": df_dep,
        "financiamento": df_fin,
        "dr": df_dr,
        "balanco": df_bal,
    }

# --- Export DOCX (Parte 2) ---
def build_docx_bp(cfg: Dict[str, Any], tabs: Dict[str, pd.DataFrame], textos: Dict[str,str], out_path: Path):
    doc = Document()
    doc.add_heading("Plano de Negócio — Parte 2", 0)

    # Secções de texto
    for titulo, txt in textos.items():
        doc.add_heading(titulo, 1)
        doc.add_paragraph(txt or "")

    # Tabelas financeiras
    def write_table(title: str, df: pd.DataFrame):
        doc.add_heading(title, 2)
        if df is None or df.empty:
            doc.add_paragraph("(sem dados)")
            return
        rows, cols = df.shape
        table = doc.add_table(rows=rows+1, cols=cols)
        table.style = "Table Grid"
        for j, col in enumerate(df.columns):
            table.cell(0, j).text = str(col)
        for i in range(rows):
            for j in range(cols):
                val = df.iat[i, j]
                table.cell(i+1, j).text = "" if pd.isna(val) else str(val)

    for nome, df in tabs.items():
        write_table(nome.upper(), df)

    doc.save(out_path)

