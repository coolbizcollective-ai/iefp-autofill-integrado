import os, json
import pandas as pd
from docx import Document
from pathlib import Path

def ia_disponivel():
    return bool(os.getenv("OPENAI_API_KEY"))

def gerar_texto_ia(titulo, instrucoes, contexto):
    if ia_disponivel():
        try:
            from openai import OpenAI
            client = OpenAI()
            prompt = f"Escreve um texto claro para '{titulo}' em PT-PT. Contexto: {json.dumps(contexto, ensure_ascii=False)}. {instrucoes}"
            resp = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role":"user","content":prompt}],
                temperature=0.5,
            )
            return resp.choices[0].message.content.strip()
        except Exception:
            return f"[Gerar com IA] {titulo}: {instrucoes}"
    return f"[Preencher] {titulo}: {instrucoes}"

def calcular_tabelas(cfg):
    """Gera tabelas exemplo para Parte 1"""
    anos = cfg.get("anos", [2025,2026,2027])
    vendas = cfg.get("vendas", [])
    df_vendas = pd.DataFrame(vendas)
    df_cogs = pd.DataFrame([{"rubrica":"COGS","valor":1000}])
    df_fse = pd.DataFrame([{"rubrica":"FSE","valor":500}])
    df_pessoal = pd.DataFrame(cfg.get("pessoal", []))
    df_dep = pd.DataFrame(cfg.get("investimento", []))
    df_fin = pd.DataFrame([{"ano":anos[0],"juros":200,"capital":500}])
    df_dr = pd.DataFrame([{"rubrica":"Resultado","valor":3000}])
    df_balanco = pd.DataFrame([{"rubrica":"Capital Próprio","valor":8000}])
    return {
        "vendas": df_vendas,
        "cogs": df_cogs,
        "fse": df_fse,
        "pessoal": df_pessoal,
        "depreciacoes": df_dep,
        "financiamento": df_fin,
        "dr": df_dr,
        "balanco": df_balanco,
    }

def build_docx(cfg, tabs, out_path: Path):
    """Gera um DOCX simples com os dados"""
    doc = Document()
    doc.add_heading("Parte 1 — Formulário IEFP", 0)

    # Identificação
    idt = cfg.get("identificacao", {})
    for k,v in idt.items():
        p = doc.add_paragraph()
        p.add_run(f"{k}: ").bold = True
        p.add_run(str(v))

    # Adiciona tabelas
    for nome, df in tabs.items():
        doc.add_heading(nome, 1)
        if not df.empty:
            table = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1])
            table.style = "Table Grid"
            for j, col in enumerate(df.columns):
                table.cell(0,j).text = str(col)
            for i in range(df.shape[0]):
                for j in range(df.shape[1]):
                    table.cell(i+1,j).text = str(df.iat[i,j])
        else:
            doc.add_paragraph("(sem dados)")

    doc.save(out_path)
